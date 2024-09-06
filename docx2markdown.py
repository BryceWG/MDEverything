from docx import Document
import re

def docx_to_markdown(input_path, output_path):
    """
    将Word文档转换为Markdown格式。

    :param input_path: Word文档的输入路径
    :param output_path: Markdown文件的输出路径
    """
    doc = Document(input_path)
    markdown_content = []

    for para in doc.paragraphs:
        markdown_content.append(process_paragraph(para))

    # 处理表格
    for table in doc.tables:
        markdown_content.extend(process_table(table))

    # 写入Markdown文件
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(markdown_content))

def process_paragraph(para):
    """处理段落,转换为Markdown格式"""
    if para.style.name.startswith('Heading'):
        level = int(para.style.name[-1])
        return '#' * level + ' ' + para.text
    else:
        text = para.text
        for run in para.runs:
            if run.bold and run.italic:
                text = re.sub(re.escape(run.text), f'***{run.text}***', text)
            elif run.bold:
                text = re.sub(re.escape(run.text), f'**{run.text}**', text)
            elif run.italic:
                text = re.sub(re.escape(run.text), f'*{run.text}*', text)
        return text

def process_table(table):
    """处理表格,转换为Markdown格式"""
    markdown_table = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.replace('\n', ' ') for cell in row.cells]
        markdown_row = '|' + '|'.join(cells) + '|'
        markdown_table.append(markdown_row)
        if i == 0:
            markdown_table.append('|' + '|'.join(['---' for _ in row.cells]) + '|')
    markdown_table.append('')  # 添加空行
    return markdown_table

# 使用示例
if __name__ == "__main__":
    input_file = "path/to/your/document.docx"
    output_file = "path/to/your/output.md"
    try:
        docx_to_markdown(input_file, output_file)
        print(f"转换成功! Markdown文件已保存为: {output_file}")
    except Exception as e:
        print(f"转换失败: {str(e)}")